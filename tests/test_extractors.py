"""Tests for Slice 4 ContentExtractor plugins.

Covers: plain, office (docx+xlsx), pdf-text, html.

Fixtures are synthesised at test-time — .docx via python-docx, .xlsx
via openpyxl, .pdf via reportlab — so no binary blobs live in the repo.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest

from common.models import (
    ChangeEvent,
    ContentRef,
    Modality,
    Operation,
)
from extractors import (
    HTMLExtractor,
    OfficeExtractor,
    PdfTextExtractor,
    PlainExtractor,
)


# ── Helpers ────────────────────────────────────────────────────────


def _event(path: str, modality: Modality = Modality.TEXT) -> ChangeEvent:
    return ChangeEvent(
        source_id="test-src",
        source_type="folder",
        operation=Operation.ADD,
        path=path,
        modality=modality,
    )


async def _make(cls):
    instance = cls(cache_dir=None, redis_client=object())  # skip cache + redis
    await instance.configure({})
    return instance


# ── plain ──────────────────────────────────────────────────────────


class TestPlainExtractor:
    async def test_inline_text_ref(self):
        ex = await _make(PlainExtractor)
        ref = ContentRef(kind="inline_text", text="hello world\n")
        ev = _event("notes.md")
        result = await ex.extract(ref, ev)

        assert result.text == "hello world\n"
        assert result.extractor_id == "plain"
        assert result.extractor_version == "0.1.0"
        assert result.metadata["format"] == "md"
        assert result.structure is None

    async def test_path_ref_json(self, tmp_path: Path):
        ex = await _make(PlainExtractor)
        p = tmp_path / "data.json"
        p.write_text('{"k": "v"}')
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("data.json")
        result = await ex.extract(ref, ev)

        assert '"k": "v"' in result.text
        assert result.metadata["format"] == "json"

    async def test_path_ref_txt(self, tmp_path: Path):
        ex = await _make(PlainExtractor)
        p = tmp_path / "doc.txt"
        p.write_text("line one\nline two\n")
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("doc.txt")
        result = await ex.extract(ref, ev)

        assert "line one" in result.text
        assert result.extractor_id == "plain"

    async def test_csv_gets_structure_hint(self, tmp_path: Path):
        ex = await _make(PlainExtractor)
        p = tmp_path / "rows.csv"
        p.write_text("a,b,c\n1,2,3\n")
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("rows.csv")
        result = await ex.extract(ref, ev)

        assert result.structure == {"format": "tabular", "delimiter": ","}

    async def test_tsv_delimiter_is_tab(self):
        ex = await _make(PlainExtractor)
        ref = ContentRef(kind="inline_text", text="a\tb\n1\t2\n")
        ev = _event("rows.tsv")
        result = await ex.extract(ref, ev)

        assert result.structure["delimiter"] == "\t"


# ── office ─────────────────────────────────────────────────────────


class TestOfficeExtractor:
    async def test_docx(self, tmp_path: Path):
        import docx
        p = tmp_path / "hello.docx"
        doc = docx.Document()
        doc.core_properties.title = "My Test Doc"
        doc.core_properties.author = "Ramón"
        doc.add_paragraph("Hello from a docx file.")
        doc.add_paragraph("Second paragraph here.")
        doc.save(str(p))

        ex = await _make(OfficeExtractor)
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("hello.docx")
        result = await ex.extract(ref, ev)

        assert "Hello from a docx file." in result.text
        assert "Second paragraph here." in result.text
        assert result.metadata.get("title") == "My Test Doc"
        assert result.metadata.get("author") == "Ramón"
        assert result.metadata["format"] == "docx"
        assert result.extractor_id == "office"
        assert result.extractor_version == "0.1.0"

    async def test_xlsx_pages_per_sheet(self, tmp_path: Path):
        import openpyxl
        p = tmp_path / "book.xlsx"
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Alpha"
        ws1.append(["h1", "h2"])
        ws1.append(["v1", "v2"])
        ws2 = wb.create_sheet("Beta")
        ws2.append(["x", "y"])
        ws2.append([1, 2])
        wb.save(str(p))

        ex = await _make(OfficeExtractor)
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("book.xlsx")
        result = await ex.extract(ref, ev)

        assert result.pages is not None
        assert len(result.pages) == 2
        titles = [pg.title for pg in result.pages]
        assert "Alpha" in titles and "Beta" in titles
        # each sheet contributed its row content
        alpha = next(pg for pg in result.pages if pg.title == "Alpha")
        assert "h1" in alpha.text and "v1" in alpha.text

    async def test_xlsx_truncated_sheets_flagged(self, tmp_path: Path):
        import openpyxl
        p = tmp_path / "big.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Big"
        for i in range(15):
            ws.append([i, i * 2])
        wb.save(str(p))

        ex = OfficeExtractor(cache_dir=None, redis_client=object())
        await ex.configure({"max_rows_per_sheet": 5})
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("big.xlsx")
        result = await ex.extract(ref, ev)

        assert result.metadata.get("truncated_sheets") == ["Big"]


# ── pdf-text ───────────────────────────────────────────────────────


def _make_simple_pdf(path: Path, text: str | None) -> None:
    """Generate a 1-page PDF. If text is None, produce a blank page."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter

    c = canvas.Canvas(str(path), pagesize=letter)
    if text:
        c.setFont("Helvetica", 12)
        c.drawString(72, 720, text)
    c.showPage()
    c.save()


class TestPdfTextExtractor:
    async def test_single_page_text(self, tmp_path: Path):
        p = tmp_path / "sample.pdf"
        _make_simple_pdf(
            p,
            "Hello from reportlab PDF body. This sentence is long enough "
            "to clearly exceed the likely-scanned threshold used by the "
            "extractor for OCR routing decisions.",
        )

        ex = await _make(PdfTextExtractor)
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("sample.pdf", modality=Modality.TEXT)
        result = await ex.extract(ref, ev)

        assert result.pages is not None
        assert len(result.pages) == 1
        assert "Hello from reportlab" in result.pages[0].text
        assert "Hello from reportlab" in result.text
        assert result.metadata["page_count"] == 1
        assert not result.metadata.get("likely_scanned")
        assert result.extractor_id == "pdf-text"
        assert result.extractor_version == "0.1.0"

    async def test_blank_pdf_flagged_as_likely_scanned(self, tmp_path: Path):
        p = tmp_path / "blank.pdf"
        _make_simple_pdf(p, None)

        ex = await _make(PdfTextExtractor)
        ref = ContentRef(kind="path", path=str(p))
        ev = _event("blank.pdf", modality=Modality.TEXT)
        result = await ex.extract(ref, ev)

        assert result.metadata.get("likely_scanned") is True
        assert result.metadata["page_count"] == 1


# ── html ───────────────────────────────────────────────────────────


_HTML_SAMPLE = """<!DOCTYPE html>
<html lang="en">
<head><title>My Article Title</title>
<meta name="author" content="Jane Reporter"/></head>
<body>
  <nav><a href="/">Home</a> | <a href="/about">About</a></nav>
  <article>
    <h1>My Article Title</h1>
    <p>This is the real article body that trafilatura should preserve.
    It contains a substantive paragraph of content roughly a few sentences
    long so the boilerplate detector keeps it. Here is another sentence.</p>
    <p>A second paragraph with more detail and context. Another sentence
    for good measure so the density heuristic picks this up cleanly.</p>
  </article>
  <footer>Copyright blah blah disclaimer navigation advert</footer>
</body>
</html>
"""


class TestHTMLExtractor:
    async def test_strips_boilerplate_and_parses_title(self):
        ex = await _make(HTMLExtractor)
        ref = ContentRef(kind="inline_text", text=_HTML_SAMPLE)
        ev = _event("article.html")
        result = await ex.extract(ref, ev)

        assert "real article body" in result.text
        # nav/footer text should be gone
        assert "Copyright blah blah" not in result.text
        assert "Home" not in result.text or "About" not in result.text
        assert result.metadata.get("title") == "My Article Title"
        assert result.extractor_id == "html"
        assert result.extractor_version == "0.1.0"
